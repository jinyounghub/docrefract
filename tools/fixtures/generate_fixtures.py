#!/usr/bin/env python3
"""Generate DocRefract's small, deterministic DOCX and PDF regression fixtures.

The fixtures are intentionally generated from source instead of being hand-authored.
Run this script with the bundled Codex Python runtime:

    python tools/fixtures/generate_fixtures.py

By default files are written to tests/DocRefract.Tests/Fixtures. Pass --out-dir
to use a scratch directory, or --check to verify committed fixtures byte-for-byte.
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import io
import json
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUT_DIR = ROOT / "tests" / "DocRefract.Tests" / "Fixtures"
FIXED_TIME = datetime(2024, 1, 2, 3, 4, 5, tzinfo=timezone.utc)
ZIP_TIME = (2024, 1, 2, 3, 4, 4)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120,
                      bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    """Apply fixed table geometry matching the standard_business_brief preset."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _configure_document(document: Document) -> None:
    """Resolve and apply the standard_business_brief fixture token map."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    props = document.core_properties
    props.author = "DocRefract"
    props.last_modified_by = "DocRefract"
    props.title = "DocRefract deterministic fixture"
    props.subject = "Regression fixture"
    props.keywords = "docrefract,fixture"
    props.comments = "Generated; do not edit by hand."
    props.created = FIXED_TIME
    props.modified = FIXED_TIME
    props.revision = 1


def _base_doc(paragraph_text: str = "Quarterly report") -> Document:
    document = Document()
    _configure_document(document)
    document.add_paragraph(paragraph_text)
    return document


def _write_docx(document: Document, path: Path) -> None:
    raw = io.BytesIO()
    document.save(raw)
    _rewrite_docx(raw.getvalue(), path)


def _rewrite_docx(data: bytes, path: Path,
                  transforms: dict[str, Callable[[bytes], bytes]] | None = None,
                  reverse_entries: bool = False) -> None:
    transforms = transforms or {}
    with zipfile.ZipFile(io.BytesIO(data), "r") as source:
        entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]
    entries.sort(key=lambda item: item[0], reverse=reverse_entries)

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED,
                         compresslevel=9, strict_timestamps=True) as target:
        for name, payload in entries:
            payload = transforms.get(name, lambda value: value)(payload)
            info = zipfile.ZipInfo(name, ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0
            target.writestr(info, payload, compress_type=zipfile.ZIP_DEFLATED,
                            compresslevel=9)
    path.write_bytes(output.getvalue())


def _metadata_transform(author: str, revision: int) -> Callable[[bytes], bytes]:
    def transform(payload: bytes) -> bytes:
        root = ET.fromstring(payload)
        author_node = root.find(f"{{{DC_NS}}}creator")
        modified_by = root.find(f"{{{CP_NS}}}lastModifiedBy")
        revision_node = root.find(f"{{{CP_NS}}}revision")
        modified = root.find(f"{{{DCTERMS_NS}}}modified")
        if author_node is not None:
            author_node.text = author
        if modified_by is not None:
            modified_by.text = author
        if revision_node is not None:
            revision_node.text = str(revision)
        if modified is not None:
            modified.text = f"2024-01-{revision + 1:02d}T03:04:05Z"
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    return transform


def _rsid_transform(rsid: str) -> Callable[[bytes], bytes]:
    def transform(payload: bytes) -> bytes:
        root = ET.fromstring(payload)
        paragraph = root.find(f".//{{{W_NS}}}p")
        run = root.find(f".//{{{W_NS}}}r")
        if paragraph is not None:
            paragraph.set(f"{{{W_NS}}}rsidR", rsid)
            paragraph.set(f"{{{W_NS}}}rsidRDefault", rsid)
        if run is not None:
            run.set(f"{{{W_NS}}}rsidRPr", rsid)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    return transform


def _style_rsid_transform(rsid: str) -> Callable[[bytes], bytes]:
    def transform(payload: bytes) -> bytes:
        root = ET.fromstring(payload)
        for node in root.findall(f".//{{{W_NS}}}rsid"):
            node.set(f"{{{W_NS}}}val", rsid)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    return transform


def _wrap_paragraph_in_content_control(paragraph, tag_value: str) -> None:
    paragraph_element = paragraph._p
    parent = paragraph_element.getparent()
    index = parent.index(paragraph_element)
    parent.remove(paragraph_element)

    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), tag_value)
    properties.append(tag)
    content = OxmlElement("w:sdtContent")
    content.append(paragraph_element)
    control.append(properties)
    control.append(content)
    parent.insert(index, control)


def _content_control_document(text: str) -> Document:
    document = Document()
    _configure_document(document)
    paragraph = document.add_paragraph(text)
    _wrap_paragraph_in_content_control(paragraph, "docrefract-fixture")
    return document


def _paragraph_sequence_document(paragraphs: list[str]) -> Document:
    document = Document()
    _configure_document(document)
    for text in paragraphs:
        document.add_paragraph(text)
    return document


def _plain_prefix_bold_tail_document(
    plain_paragraphs: list[str],
    bold_paragraphs: list[str],
) -> Document:
    document = Document()
    _configure_document(document)
    for text in plain_paragraphs:
        document.add_paragraph(text)
    for text in bold_paragraphs:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
    return document


def _bold_scope_document(bold_text: str, plain_text: str) -> Document:
    document = Document()
    _configure_document(document)
    paragraph = document.add_paragraph()
    bold = paragraph.add_run(bold_text)
    bold.bold = True
    plain = paragraph.add_run(plain_text)
    plain.bold = False
    return document


def _control_character_document(*, line_break: bool) -> Document:
    document = Document()
    _configure_document(document)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("A")
    if line_break:
        run.add_break()
    else:
        run.add_tab()
    run.add_text("B")
    return document


def _image_occurrence_document(count: int) -> Document:
    document = Document()
    _configure_document(document)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    for _ in range(count):
        run.add_picture(io.BytesIO(TINY_PNG), width=Inches(0.25))
    return document


def _image_owner_document(owner: str) -> Document:
    document = Document()
    _configure_document(document)
    paragraph_a = document.add_paragraph("A")
    paragraph_b = document.add_paragraph("B")
    target = paragraph_a if owner == "A" else paragraph_b
    target.add_run().add_picture(io.BytesIO(TINY_PNG), width=Inches(0.25))
    return document


def _floating_image_document(x_offset: int) -> Document:
    document = Document()
    _configure_document(document)
    paragraph = document.add_paragraph("Floating image")
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(TINY_PNG), width=Inches(0.25))

    drawing = run._r.find(qn("w:drawing"))
    assert drawing is not None
    inline = drawing.find(qn("wp:inline"))
    assert inline is not None

    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "0",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "column")
    horizontal_offset = OxmlElement("wp:posOffset")
    horizontal_offset.text = str(x_offset)
    horizontal.append(horizontal_offset)
    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "paragraph")
    vertical_offset = OxmlElement("wp:posOffset")
    vertical_offset.text = "0"
    vertical.append(vertical_offset)

    anchor.append(simple_position)
    anchor.append(horizontal)
    anchor.append(vertical)
    for child in list(inline):
        if child.tag == qn("wp:docPr"):
            anchor.append(OxmlElement("wp:wrapNone"))
        anchor.append(child)
    drawing.replace(inline, anchor)
    return document


def _trailing_page_break_document(has_page_break: bool) -> Document:
    document = Document()
    _configure_document(document)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Page break marker")
    if has_page_break:
        run.add_break(WD_BREAK.PAGE)
    return document


def _table_width_document(widths_dxa: list[int]) -> Document:
    document = Document()
    _configure_document(document)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Left"
    table.cell(0, 1).text = "Right"
    _set_table_geometry(table, widths_dxa)
    return document


def _create_docx_fixtures(out_dir: Path) -> None:
    metadata_doc = _base_doc("The semantic content is unchanged.")
    raw = io.BytesIO()
    metadata_doc.save(raw)
    common = raw.getvalue()

    _rewrite_docx(
        common,
        out_dir / "docx_metadata_before.docx",
        {
            "docProps/core.xml": _metadata_transform("Before Author", 1),
            "word/document.xml": _rsid_transform("00112233"),
        },
    )
    _rewrite_docx(
        common,
        out_dir / "docx_metadata_after.docx",
        {
            "docProps/core.xml": _metadata_transform("After Author", 7),
            "word/document.xml": _rsid_transform("AABBCCDD"),
        },
        reverse_entries=True,
    )

    _write_docx(
        _base_doc("Quarterly revenue is 120 USD."),
        out_dir / "docx_text_before.docx",
    )
    _write_docx(
        _base_doc("Quarterly revenue is 125 USD."),
        out_dir / "docx_text_after.docx",
    )

    for name, value in (("before", "Pending"), ("after", "Approved")):
        document = _base_doc("Release matrix")
        table = document.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Status"
        table.cell(1, 0).text = "Desktop CLI"
        table.cell(1, 1).text = value
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        _set_table_geometry(table, [4680, 4680])
        _write_docx(document, out_dir / f"docx_table_{name}.docx")

    for name, bold in (("before", False), ("after", True)):
        document = Document()
        _configure_document(document)
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Formatting-only signal")
        run.bold = bold
        _write_docx(document, out_dir / f"docx_style_{name}.docx")

    _write_docx(
        _content_control_document("Content control value: before"),
        out_dir / "docx_sdt_before.docx",
    )
    _write_docx(
        _content_control_document("Content control value: after"),
        out_dir / "docx_sdt_after.docx",
    )

    stable_paragraphs = [f"Stable paragraph {index:04d}" for index in range(1, 2051)]
    inserted_paragraphs = [f"Inserted paragraph {index:04d}" for index in range(1, 41)]
    _write_docx(
        _paragraph_sequence_document(stable_paragraphs),
        out_dir / "docx_large_insert_before.docx",
    )
    _write_docx(
        _paragraph_sequence_document(inserted_paragraphs + stable_paragraphs),
        out_dir / "docx_large_insert_after.docx",
    )

    _write_docx(
        _paragraph_sequence_document(stable_paragraphs),
        out_dir / "docx_large_insert_format_before.docx",
    )
    _write_docx(
        _plain_prefix_bold_tail_document(inserted_paragraphs, stable_paragraphs),
        out_dir / "docx_large_insert_format_after.docx",
    )

    _write_docx(
        _paragraph_sequence_document(["A", "B", "C"]),
        out_dir / "docx_move_before.docx",
    )
    _write_docx(
        _paragraph_sequence_document(["B", "C", "A"]),
        out_dir / "docx_move_after.docx",
    )

    _write_docx(
        _bold_scope_document("A", "BC"),
        out_dir / "docx_bold_scope_before.docx",
    )
    _write_docx(
        _bold_scope_document("AB", "C"),
        out_dir / "docx_bold_scope_after.docx",
    )

    style_rsid_document = _base_doc("Style rsid changes are metadata only.")
    style_rsid_document.paragraphs[0].style = "Heading 1"
    raw = io.BytesIO()
    style_rsid_document.save(raw)
    style_rsid_common = raw.getvalue()
    _rewrite_docx(
        style_rsid_common,
        out_dir / "docx_style_rsid_before.docx",
        {"word/styles.xml": _style_rsid_transform("00112233")},
    )
    _rewrite_docx(
        style_rsid_common,
        out_dir / "docx_style_rsid_after.docx",
        {"word/styles.xml": _style_rsid_transform("AABBCCDD")},
    )

    _write_docx(
        _control_character_document(line_break=False),
        out_dir / "docx_tab_before.docx",
    )
    _write_docx(
        _control_character_document(line_break=True),
        out_dir / "docx_tab_after.docx",
    )

    _write_docx(
        _image_occurrence_document(1),
        out_dir / "docx_image_occurrence_before.docx",
    )
    _write_docx(
        _image_occurrence_document(2),
        out_dir / "docx_image_occurrence_after.docx",
    )

    _write_docx(
        _image_owner_document("A"),
        out_dir / "docx_image_move_before.docx",
    )
    _write_docx(
        _image_owner_document("B"),
        out_dir / "docx_image_move_after.docx",
    )

    _write_docx(
        _floating_image_document(0),
        out_dir / "docx_floating_image_position_before.docx",
    )
    _write_docx(
        _floating_image_document(914400),
        out_dir / "docx_floating_image_position_after.docx",
    )

    _write_docx(
        _trailing_page_break_document(False),
        out_dir / "docx_page_break_before.docx",
    )
    _write_docx(
        _trailing_page_break_document(True),
        out_dir / "docx_page_break_after.docx",
    )

    _write_docx(
        _table_width_document([4680, 4680]),
        out_dir / "docx_table_width_before.docx",
    )
    _write_docx(
        _table_width_document([3600, 5760]),
        out_dir / "docx_table_width_after.docx",
    )


def _write_pdf(
    path: Path,
    text: str,
    *, fill_rgb: tuple[float, float, float] = (0, 0, 0),
) -> None:
    # invariant=1 removes creation timestamps and stabilizes the trailer ID.
    pdf = canvas.Canvas(
        str(path),
        pagesize=(612, 792),
        bottomup=1,
        pageCompression=0,
        invariant=1,
    )
    pdf.setTitle("DocRefract deterministic fixture")
    pdf.setAuthor("DocRefract")
    pdf.setFont("Helvetica", 12)
    pdf.setFillColorRGB(*fill_rgb)
    pdf.drawString(72, 720, text)
    pdf.showPage()
    pdf.save()


def _create_pdf_fixtures(out_dir: Path) -> None:
    _write_pdf(out_dir / "pdf_text_before.pdf", "Invoice total: 120 USD")
    _write_pdf(out_dir / "pdf_text_after.pdf", "Invoice total: 125 USD")
    _write_pdf(out_dir / "pdf_identical.pdf", "No changes on this page")
    _write_pdf(
        out_dir / "pdf_color_before.pdf",
        "Color-only signal",
        fill_rgb=(0, 0, 0),
    )
    _write_pdf(
        out_dir / "pdf_color_after.pdf",
        "Color-only signal",
        fill_rgb=(1, 0, 0),
    )


def _manifest(out_dir: Path) -> dict[str, object]:
    records: list[dict[str, object]] = []
    for path in sorted(out_dir.glob("*")):
        if path.name == "manifest.json" or not path.is_file():
            continue
        data = path.read_bytes()
        records.append(
            {
                "name": path.name,
                "bytes": len(data),
                "sha256": hashlib.sha256(data).hexdigest(),
            }
        )
    return {"schemaVersion": 1, "files": records}


def _manifest_bytes(out_dir: Path) -> bytes:
    text = json.dumps(
        _manifest(out_dir),
        ensure_ascii=False,
        indent=2,
        sort_keys=False,
    )
    return (text + "\n").encode("utf-8")


def _generate(out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    for path in out_dir.iterdir():
        if path.is_file():
            path.unlink()
    _create_docx_fixtures(out_dir)
    _create_pdf_fixtures(out_dir)
    (out_dir / "manifest.json").write_bytes(_manifest_bytes(out_dir))


def _check(expected_dir: Path) -> int:
    if not expected_dir.exists():
        print(f"Fixture directory does not exist: {expected_dir}", file=sys.stderr)
        return 1
    # Keep scratch data inside the writable project tree. A stable directory is
    # used because some Windows restricted tokens cannot re-open directories
    # created with tempfile's owner-only ACL.
    actual_dir = Path(__file__).parent / ".check"
    actual_dir.mkdir(exist_ok=True)
    try:
        _generate(actual_dir)
        expected_names = sorted(path.name for path in expected_dir.iterdir() if path.is_file())
        actual_names = sorted(path.name for path in actual_dir.iterdir() if path.is_file())
        if expected_names != actual_names:
            print("Fixture filenames differ.", file=sys.stderr)
            print(f"expected: {expected_names}", file=sys.stderr)
            print(f"actual:   {actual_names}", file=sys.stderr)
            return 1
        mismatches = [
            name
            for name in expected_names
            if (expected_dir / name).read_bytes() != (actual_dir / name).read_bytes()
        ]
        if mismatches:
            print(f"Non-deterministic or stale fixtures: {', '.join(mismatches)}",
                  file=sys.stderr)
            return 1
    finally:
        for path in actual_dir.iterdir():
            if path.is_file():
                path.unlink()
        actual_dir.rmdir()
    print(f"Fixtures are reproducible: {expected_dir}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=DEFAULT_OUT_DIR,
        help=f"output directory (default: {DEFAULT_OUT_DIR})",
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="regenerate in a temporary directory and compare byte-for-byte",
    )
    args = parser.parse_args()
    out_dir = args.out_dir.resolve()
    if args.check:
        return _check(out_dir)
    _generate(out_dir)
    print(f"Generated {len(_manifest(out_dir)['files'])} fixtures in {out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
